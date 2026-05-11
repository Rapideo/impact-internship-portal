"""Build the IMPACT prototype walkthrough handout (.docx)."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve().parent
SCREENS = HERE / "screens"
OUT = HERE / "IMPACT_Prototype_Walkthrough.docx"

# Brand palette
NAVY = RGBColor(0x15, 0x3A, 0x98)
NAVY_DEEP = RGBColor(0x05, 0x10, 0x28)
GOLD = RGBColor(0xFF, 0xD7, 0x1F)
INK = RGBColor(0x1A, 0x22, 0x35)
MUTED = RGBColor(0x55, 0x5E, 0x70)


# ----- low-level helpers -----

def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, sides=("top", "bottom", "left", "right"),
                     size="6", color="D6DAE3"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side in sides:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def remove_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side in ("top", "bottom", "left", "right", "insideH", "insideV"):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'nil')
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    """Cell padding in twips (1/20 pt)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('left', left),
                      ('bottom', bottom), ('right', right)):
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tc_pr.append(tcMar)


def add_run(p, text, *, bold=False, color=None, size=None, font="Calibri",
            all_caps=False, italic=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = font
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if all_caps:
        rPr = r._element.get_or_add_rPr()
        caps = OxmlElement('w:caps')
        caps.set(qn('w:val'), '1')
        rPr.append(caps)
    return r


def add_para(doc_or_cell, text="", *, bold=False, color=None, size=11,
             font="Calibri", align=None, all_caps=False, space_before=0,
             space_after=4, italic=False):
    p = doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        add_run(p, text, bold=bold, color=color, size=size, font=font,
                all_caps=all_caps, italic=italic)
    return p


def hr_paragraph(doc, color="153A98", size="12"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), size)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)


# ----- composite layout helpers -----

def title_block(doc):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.8)
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.8)
    set_cell_bg(cell, "051028")
    set_cell_margins(cell, top=140, bottom=140, left=240, right=240)
    remove_cell_borders(cell)
    cell.text = ""

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "IMPACT  ·  INTERNSHIP ASSESSMENT PORTAL",
            bold=True, color=GOLD, size=8.5, all_caps=True)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(0)
    add_run(p2, "Prototype Walkthrough",
            bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=18)

    p3 = cell.add_paragraph()
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after = Pt(0)
    add_run(p3,
            "A guided tour of the intern and administrator experiences "
            "in the IMPACT clickable prototype.",
            color=RGBColor(0xC9, 0xD2, 0xE0), size=9.5)

    p4 = cell.add_paragraph()
    p4.paragraph_format.space_before = Pt(6)
    p4.paragraph_format.space_after = Pt(0)
    add_run(p4, "TRY IT LIVE  →  ",
            bold=True, color=GOLD, size=8.5, all_caps=True)
    add_run(p4, "impact-internship-portal.netlify.app",
            bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10.5)


def section_header(doc, eyebrow, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, eyebrow, bold=True, color=GOLD, size=8, all_caps=True)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    add_run(p2, title, bold=True, color=NAVY_DEEP, size=14)
    hr_paragraph(doc, color="153A98", size="12")


def step_row(doc, num, step_title, narrative, image_filename,
             image_width_in=2.7):
    """Two-column row: image left, narrative right."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Inches(2.85)
    tbl.columns[1].width = Inches(3.95)
    tbl.allow_autofit = False

    img_cell, txt_cell = tbl.rows[0].cells
    img_cell.width = Inches(2.85)
    txt_cell.width = Inches(3.95)
    img_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    txt_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    remove_cell_borders(img_cell)
    remove_cell_borders(txt_cell)
    set_cell_margins(img_cell, top=10, bottom=10, left=0, right=100)
    set_cell_margins(txt_cell, top=10, bottom=10, left=60, right=20)

    img_p = img_cell.paragraphs[0]
    img_p.paragraph_format.space_after = Pt(0)
    img_run = img_p.add_run()
    img_run.add_picture(str(SCREENS / image_filename),
                        width=Inches(image_width_in))

    # Narrative side: step number + title + body
    head = txt_cell.paragraphs[0]
    head.paragraph_format.space_after = Pt(0)
    add_run(head, f"STEP {num:02d}",
            bold=True, color=GOLD, size=8.5, all_caps=True)

    title_p = txt_cell.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    add_run(title_p, step_title, bold=True, color=NAVY_DEEP, size=11.5)

    body_p = txt_cell.add_paragraph()
    body_p.paragraph_format.space_before = Pt(0)
    body_p.paragraph_format.space_after = Pt(0)
    body_p.paragraph_format.line_spacing = 1.15
    add_run(body_p, narrative, color=INK, size=9.5)

    # Trailing tiny gap below the row
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    # Make spacer line-height very small
    spacer.paragraph_format.line_spacing = 1.0
    for r in spacer.runs:
        r.font.size = Pt(2)


def callout(doc, label, body):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.8)
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.8)
    set_cell_bg(cell, "FFF8DA")
    set_cell_margins(cell, top=110, bottom=110, left=160, right=160)
    set_cell_borders(cell, sides=("left",), size="36", color="FFD71F")
    set_cell_borders(cell, sides=("top", "bottom", "right"),
                     size="4", color="EFE3A8")

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    add_run(p, label, bold=True, color=NAVY_DEEP, size=8.5, all_caps=True)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.18
    add_run(p2, body, color=INK, size=9.5)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ----- document body -----

def build():
    doc = Document()

    # Page setup — narrow-ish margins for a denser, magazine feel.
    section = doc.sections[0]
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

    # Default style
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK

    # ---- Title bar ----
    title_block(doc)

    # =====================================================================
    # SECTION 1 — INTERN EXPERIENCE
    # =====================================================================
    section_header(doc, "Audience 01", "The Intern Experience")

    lead = doc.add_paragraph()
    lead.paragraph_format.space_before = Pt(0)
    lead.paragraph_format.space_after = Pt(4)
    lead.paragraph_format.line_spacing = 1.18
    add_run(lead,
            "Interns never create an account — they visit a public link, "
            "pick the reflection that matches where they are in the "
            "program, and submit. We walk through ",
            color=INK, size=9.5)
    add_run(lead, "Personal Goals", bold=True, color=NAVY_DEEP, size=9.5)
    add_run(lead,
            " below; Midpoint Reflection and Participant Feedback follow "
            "the exact same four-step pattern.",
            color=INK, size=9.5)

    step_row(doc, 1, "Land on the public homepage",
             "A branded landing page introduces the program and offers two "
             "primary calls-to-action. Both buttons route to the assessment "
             "chooser — interns don’t have to know which form they need yet.",
             "01-landing.png")

    step_row(doc, 2, "Choose the right reflection",
             "Three cards label each reflection by program stage: At Start, "
             "At Midpoint, At Exit. Cards are status-aware — once an "
             "assessment is submitted, its card switches to a completed state "
             "to prevent duplicate submissions.",
             "02-chooser.png")

    step_row(doc, 3, "Identify yourself, then reflect",
             "The intern enters Last Name + Cohort + Zipcode (the unique "
             "identifier used throughout the system, with no password to "
             "manage) and answers the reflection questions. Cohort options "
             "are pulled from the live cohort list maintained by admins.",
             "03-personal-goals.png")

    step_row(doc, 4, "Confirmation — once and done",
             "On submit, the intern lands on a thank-you page. Each "
             "reflection can only be submitted once per intern; that "
             "completion is now visible to program staff inside the admin "
             "portal alongside the rest of that intern’s record.",
             "04-confirmation.png")

    callout(doc, "The other two reflections work identically",
            "Midpoint Reflection (8 questions) and Participant Feedback "
            "(7 questions, mixed format including Likert and Yes/No) "
            "follow the exact same four-step flow. The chooser hub is the "
            "single entry point for all three.")

    # =====================================================================
    # SECTION 2 — ADMIN EXPERIENCE  (flows naturally onto next page)
    # =====================================================================
    section_header(doc, "Audience 02", "The Administrator Experience")

    lead = doc.add_paragraph()
    lead.paragraph_format.space_before = Pt(0)
    lead.paragraph_format.space_after = Pt(4)
    lead.paragraph_format.line_spacing = 1.18
    add_run(lead,
            "Program staff sign in once and manage the full cohort "
            "lifecycle from a single navigation. Below is the most common "
            "day-to-day path — taking an intern from intake to outcomes.",
            color=INK, size=9.5)

    step_row(doc, 1, "Land on a program-level dashboard",
             "After sign-in, the admin sees a personalized greeting plus "
             "three live KPIs — active cohorts, active interns, and "
             "confirmed 90-day outcomes. Quick-link buttons drop them "
             "straight into the four most-used areas.",
             "05-admin-home.png")

    step_row(doc, 2, "Browse and filter the intern roster",
             "The Interns page is the working list: searchable by name or "
             "cohort, filterable by outcome status, and sortable by start "
             "date. Each row shows current phase and 90-day status at a "
             "glance, with one click into the full record.",
             "06-interns.png")

    step_row(doc, 3, "Manage one intern, one record",
             "The unified intern record holds everything for that "
             "participant — personal info, internship assignment, the "
             "Entry Assessment barriers checklist, self-assessments, "
             "competency evaluations, and 90/180-day outcome tracking — "
             "in six numbered panels on a single page.",
             "07-intern-record.png")

    step_row(doc, 4, "Launch the right assessment",
             "From the Assessments hub, staff pick the assessment to "
             "administer (Competency or Exit Employer Survey), then choose "
             "the intern. The form opens pre-filled with that intern’s "
             "identity, ready for a phase-specific evaluation.",
             "08-assessments-hub.png")

    step_row(doc, 5, "Score against the role-specific rubric",
             "Competency questions stitch three layers automatically: a "
             "shared core of professional competencies, role-specific "
             "skills for the intern’s cohort, and any per-intern customizations "
             "— so staff always score against the right rubric without "
             "managing it manually.",
             "09-competency.png")

    step_row(doc, 6, "Read the program-level story",
             "Reports aggregates the same data across the cohort cycle: "
             "top entry barriers, competency progression by phase, and "
             "placement outcomes. Charts in this prototype use demo data "
             "and will be replaced with live queries in the production build.",
             "10-reports.png")

    callout(doc, "What this prototype is — and isn’t",
            "Every page above is a clickable static mockup on a shared "
            "demo dataset; submissions persist for the browser session so "
            "stakeholders can experience full flows. The next phase converts "
            "this front-end into a production app with a real database, "
            "authentication, and live reporting.")

    doc.save(str(OUT))
    print(f"wrote {OUT}  ({OUT.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    build()
